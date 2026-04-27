"""Build a Word document installation guide (Arabic, RTL)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def set_rtl(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_rtl(cell):
    for p in cell.paragraphs:
        set_rtl(p)


def add_heading_ar(doc, text, level=1, color=None):
    h = doc.add_heading(level=level)
    set_rtl(h)
    run = h.add_run(text)
    run.font.name = "Calibri"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:cs"), "Arial")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    if color:
        run.font.color.rgb = color
    return h


def add_para_ar(doc, text, bold=False, size=11, color=None, style=None):
    p = doc.add_paragraph(style=style)
    set_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:cs"), "Arial")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet_ar(doc, text):
    return add_para_ar(doc, text, style="List Bullet")


def add_numbered_ar(doc, text):
    return add_para_ar(doc, text, style="List Number")


def add_code_block(doc, code, lang_label=None):
    if lang_label:
        p = doc.add_paragraph()
        set_rtl(p)
        r = p.add_run(lang_label)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = table.rows[0].cells[0]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F3F4F6")
    cell._tc.get_or_add_tcPr().append(shading)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)
    doc.add_paragraph()


def add_table_ar(doc, headers, rows, header_color="2563EB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        set_rtl(p)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), header_color)
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            set_rtl(p)
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()


def add_callout(doc, title, body, color="DBEAFE", border="2563EB"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)
    cell.text = ""
    p1 = cell.paragraphs[0]
    set_rtl(p1)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p2 = cell.add_paragraph()
    set_rtl(p2)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        sectPr = section._sectPr
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)

    # Default font for normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===== Cover =====
    title = doc.add_paragraph()
    set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("دليل تركيب وتشغيل\nنظام الفاتورة الإلكترونية السعودية (ZATCA)")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    sub = doc.add_paragraph()
    set_rtl(sub)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("التركيب على سيرفر محلي وعلى سيرفر سحابي")
    rs.font.size = Pt(14)
    rs.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()
    add_para_ar(
        doc,
        "هذا الدليل يشرح خطوة بخطوة كيفية تجهيز البيئة، تركيب الاعتماديات، إعداد قاعدة البيانات، تشغيل التطبيق في وضع التطوير، ثم نشره على سيرفر إنتاجي سحابي مع الحماية بـ HTTPS وعكس البروكسي عبر Nginx وإدارة العمليات بـ PM2.",
        size=11,
    )

    doc.add_page_break()

    # ===== Table of contents =====
    add_heading_ar(doc, "المحتويات", level=1, color=RGBColor(0x1E, 0x3A, 0x8A))
    toc = [
        "1. نظرة عامة على بنية المشروع",
        "2. المتطلبات الأساسية",
        "3. التركيب على سيرفر محلي (للتطوير أو الاستخدام الداخلي)",
        "4. التركيب على سيرفر سحابي للإنتاج",
        "5. متغيرات البيئة (الشرح الكامل)",
        "6. التحديثات وعمليات الصيانة",
        "7. حل المشكلات الشائعة",
    ]
    for item in toc:
        add_para_ar(doc, item, size=11)
    doc.add_page_break()

    # ===== 1) Overview =====
    add_heading_ar(doc, "1. نظرة عامة على بنية المشروع", level=1, color=RGBColor(0x1E, 0x3A, 0x8A))
    add_para_ar(
        doc,
        "المشروع مستودع موحّد (monorepo) بإدارة pnpm، يضم عدة تطبيقات ومكتبات مشتركة:",
    )
    add_bullet_ar(doc, "artifacts/zatca-invoicing — الواجهة الرئيسية بالعربية (React + Vite). تحوي شاشات الفواتير والعملاء والمخزون والموارد البشرية والمحاسبة وأدوات الذكاء الاصطناعي وإعدادات السوبر أدمن.")
    add_bullet_ar(doc, "artifacts/api-server — خادم الـ API (Node.js + Express). يحتوي المسارات، وسطاء فحص الصلاحيات، خدمات ZATCA للتوقيع والإرسال، ومحرّك أرقام التسلسل.")
    add_bullet_ar(doc, "artifacts/pos — تطبيق نقاط البيع منفصل بواجهته الخاصة.")
    add_bullet_ar(doc, "artifacts/mockup-sandbox — بيئة معاينة المكوّنات للتصميم.")
    add_bullet_ar(doc, "lib/ — مكتبات مشتركة (مخططات Drizzle، أنواع TypeScript، مخطط OpenAPI).")
    add_bullet_ar(doc, "scripts/ — سكربتات بناء وصيانة.")
    add_para_ar(
        doc,
        "الكل يعتمد على قاعدة بيانات PostgreSQL واحدة، وتخزين ملفات (محلي أو على سحابة).",
    )

    doc.add_page_break()

    # ===== 2) Prerequisites =====
    add_heading_ar(doc, "2. المتطلبات الأساسية", level=1, color=RGBColor(0x1E, 0x3A, 0x8A))
    add_table_ar(
        doc,
        ["المتطلب", "الإصدار الموصى به", "السبب"],
        [
            ["Node.js", "20.x أو أحدث", "بيئة تشغيل الواجهة والخادم"],
            ["pnpm", "9.x أو أحدث", "مدير الحزم (المشروع monorepo)"],
            ["PostgreSQL", "14 أو أحدث", "قاعدة البيانات الرئيسية"],
            ["Git", "أي حديث", "استنساخ وتحديث المشروع"],
            ["Nginx", "1.20+ (للإنتاج)", "عكس البروكسي وشهادات HTTPS"],
            ["PM2", "5.x (للإنتاج)", "إدارة عمليات Node وإعادة تشغيلها تلقائيًا"],
        ],
    )

    add_heading_ar(doc, "تركيب المتطلبات على Ubuntu/Debian", level=2)
    add_code_block(doc, """# تحديث الحزم
sudo apt update && sudo apt upgrade -y

# Node.js 20
curl -fsSL https://deb.nodesource.com/setup_20.x | sudo -E bash -
sudo apt install -y nodejs

# pnpm
sudo npm install -g pnpm

# PostgreSQL
sudo apt install -y postgresql postgresql-contrib

# أدوات الإنتاج (للسيرفر السحابي فقط)
sudo apt install -y nginx
sudo npm install -g pm2

# Git
sudo apt install -y git""", "Bash")

    add_heading_ar(doc, "تركيب المتطلبات على ويندوز", level=2)
    add_bullet_ar(doc, "Node.js: نزّل من nodejs.org النسخة LTS وثبّتها.")
    add_bullet_ar(doc, "pnpm: افتح PowerShell كمدير ونفّذ: npm install -g pnpm")
    add_bullet_ar(doc, "PostgreSQL: نزّل من postgresql.org/download/windows")
    add_bullet_ar(doc, "Git: نزّل من git-scm.com")

    add_heading_ar(doc, "تركيب المتطلبات على ماك", level=2)
    add_code_block(doc, """brew install node@20 pnpm postgresql@16 git
brew services start postgresql@16""", "Bash")

    doc.add_page_break()

    # ===== 3) Local install =====
    add_heading_ar(doc, "3. التركيب على سيرفر محلي", level=1, color=RGBColor(0x05, 0x96, 0x69))
    add_para_ar(doc, "هذا القسم يناسب التطوير على جهازك أو التشغيل داخل شبكة الشركة.", bold=True)

    add_heading_ar(doc, "الخطوة 1 — فك ضغط المشروع", level=2)
    add_code_block(doc, """tar -xzf zatca-source-20260427-231316.tar.gz -C ~/zatca
cd ~/zatca""", "Bash")

    add_heading_ar(doc, "الخطوة 2 — تركيب الاعتماديات", level=2)
    add_code_block(doc, "pnpm install", "Bash")
    add_para_ar(doc, "هذه العملية قد تستغرق 3-5 دقائق في أول مرة (تنزّل ~600 ميجا من الحزم).")

    add_heading_ar(doc, "الخطوة 3 — إنشاء قاعدة البيانات", level=2)
    add_code_block(doc, """sudo -u postgres psql

CREATE DATABASE zatca;
CREATE USER zatca_user WITH PASSWORD 'كلمة_مرور_قوية_هنا';
GRANT ALL PRIVILEGES ON DATABASE zatca TO zatca_user;
ALTER DATABASE zatca OWNER TO zatca_user;
\\q""", "SQL")

    add_heading_ar(doc, "الخطوة 4 — إنشاء ملف متغيرات البيئة", level=2)
    add_para_ar(doc, "أنشئ ملف .env في جذر المشروع:")
    add_code_block(doc, """# قاعدة البيانات
DATABASE_URL=postgresql://zatca_user:كلمة_المرور@localhost:5432/zatca

# الجلسات (سلسلة عشوائية ≥ 64 حرف)
SESSION_SECRET=ضع_هنا_سلسلة_طويلة_عشوائية_على_الأقل_أربعة_وستين_حرفًا

# تخزين الملفات (مجلد محلي)
DEFAULT_OBJECT_STORAGE_BUCKET_ID=local-bucket
PRIVATE_OBJECT_DIR=/var/zatca/storage/private
PUBLIC_OBJECT_SEARCH_PATHS=/var/zatca/storage/public

# اختيارية — البريد (للإشعارات والتقارير)
SMTP_HOST=smtp.gmail.com
SMTP_PORT=587
SMTP_USER=your-email@gmail.com
SMTP_PASS=app-password-here
SMTP_FROM=noreply@yourcompany.sa

# اختيارية — التحقق من العنوان السعودي
SPL_API_KEY=مفتاح_من_البريد_السعودي

# اختيارية — حماية النماذج من البوتات (Cloudflare Turnstile)
TURNSTILE_SITE_KEY=...
TURNSTILE_SECRET_KEY=...
VITE_TURNSTILE_SITE_KEY=...""", ".env")

    add_para_ar(doc, "ثم أنشئ مجلدات التخزين:")
    add_code_block(doc, """sudo mkdir -p /var/zatca/storage/{private,public}
sudo chown -R $USER:$USER /var/zatca""", "Bash")

    add_heading_ar(doc, "الخطوة 5 — دفع مخطط قاعدة البيانات", level=2)
    add_code_block(doc, "pnpm --filter @workspace/api-server run db:push", "Bash")
    add_para_ar(doc, "هذا الأمر ينشئ تلقائيًا 126 جدولًا و26 نوع enum — قد يطلب التأكيد عند تغييرات قد تؤثر على البيانات؛ في أول تركيب وافق على الكل.")

    add_heading_ar(doc, "الخطوة 6 — تشغيل التطبيقات في وضع التطوير", level=2)
    add_para_ar(doc, "افتح ثلاث نوافذ طرفية منفصلة:")
    add_code_block(doc, """# نافذة 1 — خادم الـ API (يستمع على 8080)
pnpm --filter @workspace/api-server run dev

# نافذة 2 — الواجهة الرئيسية (تستمع على 5173)
pnpm --filter @workspace/zatca-invoicing run dev

# نافذة 3 — نقاط البيع (اختياري)
pnpm --filter @workspace/pos run dev""", "Bash")

    add_para_ar(doc, "افتح المتصفح على:")
    add_bullet_ar(doc, "الواجهة الرئيسية: http://localhost:5173")
    add_bullet_ar(doc, "خادم الـ API: http://localhost:8080")
    add_bullet_ar(doc, "نقاط البيع: http://localhost:5174")

    add_heading_ar(doc, "الخطوة 7 — إنشاء حساب السوبر أدمن", level=2)
    add_para_ar(doc, "أول مستخدم يُسجَّل من شاشة التسجيل يصبح سوبر أدمن تلقائيًا. للترقية اليدوية لاحقًا:")
    add_code_block(doc, """sudo -u postgres psql -d zatca -c "UPDATE users SET role='superadmin' WHERE email='you@email.com';" """, "SQL")

    add_callout(
        doc,
        "ملاحظة:",
        "في وضع التطوير، الواجهة تتصل بالخادم تلقائيًا عبر منفذ 8080. لو غيّرت المنافذ، حدّث ملف vite.config.ts و .env بما يناسب.",
    )

    doc.add_page_break()

    # ===== 4) Cloud server =====
    add_heading_ar(doc, "4. التركيب على سيرفر سحابي للإنتاج", level=1, color=RGBColor(0xDC, 0x26, 0x26))
    add_para_ar(
        doc,
        "هذا القسم يشرح نشر التطبيق على VPS (مثل DigitalOcean, AWS EC2, Hetzner, Linode, Contabo, أو أي سيرفر Ubuntu 22.04+).",
        bold=True,
    )

    add_heading_ar(doc, "المواصفات الموصى بها للسيرفر", level=2)
    add_table_ar(
        doc,
        ["العنصر", "الحد الأدنى", "موصى به للإنتاج"],
        [
            ["المعالج", "2 vCPU", "4 vCPU"],
            ["الذاكرة", "4 GB RAM", "8 GB RAM"],
            ["التخزين", "40 GB SSD", "100 GB SSD"],
            ["نظام التشغيل", "Ubuntu 22.04 LTS", "Ubuntu 22.04 / 24.04 LTS"],
            ["الشبكة", "IPv4 ثابت", "IPv4 + اسم نطاق مربوط"],
        ],
    )

    add_heading_ar(doc, "الخطوة 1 — تجهيز السيرفر وتأمينه", level=2)
    add_code_block(doc, """# الاتصال بالسيرفر
ssh root@YOUR_SERVER_IP

# إنشاء مستخدم بصلاحيات sudo (لا تستخدم root مباشرة)
adduser zatca
usermod -aG sudo zatca

# تفعيل الجدار الناري
ufw allow OpenSSH
ufw allow 'Nginx Full'
ufw enable

# الانتقال للمستخدم الجديد
su - zatca""", "Bash")

    add_heading_ar(doc, "الخطوة 2 — تركيب المتطلبات", level=2)
    add_code_block(doc, """sudo apt update && sudo apt upgrade -y

# Node.js 20
curl -fsSL https://deb.nodesource.com/setup_20.x | sudo -E bash -
sudo apt install -y nodejs

# pnpm + PM2
sudo npm install -g pnpm pm2

# PostgreSQL + Nginx + Certbot لشهادات HTTPS
sudo apt install -y postgresql postgresql-contrib nginx certbot python3-certbot-nginx git""", "Bash")

    add_heading_ar(doc, "الخطوة 3 — رفع المشروع إلى السيرفر", level=2)
    add_para_ar(doc, "خياران:")
    add_para_ar(doc, "أ) عبر Git (الأفضل):", bold=True)
    add_code_block(doc, """cd ~
git clone https://github.com/your-org/zatca.git
cd zatca""", "Bash")
    add_para_ar(doc, "ب) عبر رفع الأرشيف:", bold=True)
    add_code_block(doc, """# من جهازك المحلي
scp zatca-source-20260427-231316.tar.gz zatca@YOUR_SERVER_IP:~/

# على السيرفر
mkdir -p ~/zatca && tar -xzf ~/zatca-source-20260427-231316.tar.gz -C ~/zatca
cd ~/zatca""", "Bash")

    add_heading_ar(doc, "الخطوة 4 — تركيب الاعتماديات وبناء الإنتاج", level=2)
    add_code_block(doc, """pnpm install --frozen-lockfile

# بناء الواجهات للإنتاج
pnpm --filter @workspace/zatca-invoicing run build
pnpm --filter @workspace/pos run build
pnpm --filter @workspace/api-server run build""", "Bash")

    add_heading_ar(doc, "الخطوة 5 — إعداد قاعدة البيانات الإنتاجية", level=2)
    add_code_block(doc, """sudo -u postgres psql

CREATE DATABASE zatca_prod;
CREATE USER zatca_prod_user WITH PASSWORD 'كلمة_مرور_قوية_جدًا_عشوائية';
GRANT ALL PRIVILEGES ON DATABASE zatca_prod TO zatca_prod_user;
ALTER DATABASE zatca_prod OWNER TO zatca_prod_user;
\\q""", "SQL")

    add_heading_ar(doc, "الخطوة 6 — متغيرات البيئة الإنتاجية", level=2)
    add_para_ar(doc, "أنشئ ملف .env في جذر المشروع:")
    add_code_block(doc, """NODE_ENV=production
PORT=8080

# قاعدة البيانات
DATABASE_URL=postgresql://zatca_prod_user:PASSWORD@localhost:5432/zatca_prod

# الجلسات
SESSION_SECRET=سلسلة_عشوائية_طويلة_جدًا_للإنتاج_فقط

# التخزين
DEFAULT_OBJECT_STORAGE_BUCKET_ID=production-bucket
PRIVATE_OBJECT_DIR=/var/zatca/storage/private
PUBLIC_OBJECT_SEARCH_PATHS=/var/zatca/storage/public

# البريد
SMTP_HOST=smtp.your-provider.com
SMTP_PORT=587
SMTP_USER=...
SMTP_PASS=...
SMTP_FROM=invoices@yourcompany.sa

# مفاتيح اختيارية
SPL_API_KEY=...
TURNSTILE_SECRET_KEY=...
TURNSTILE_SITE_KEY=...
VITE_TURNSTILE_SITE_KEY=...""", ".env")

    add_callout(
        doc,
        "تنبيه أمني هام:",
        "ضع صلاحيات صارمة على ملف .env: chmod 600 .env — لمنع المستخدمين الآخرين من قراءته.",
        color="FEE2E2",
        border="DC2626",
    )

    add_heading_ar(doc, "الخطوة 7 — دفع مخطط القاعدة", level=2)
    add_code_block(doc, "pnpm --filter @workspace/api-server run db:push", "Bash")

    add_heading_ar(doc, "الخطوة 8 — تشغيل الخدمات بـ PM2", level=2)
    add_para_ar(doc, "أنشئ ملف ecosystem.config.cjs في جذر المشروع:")
    add_code_block(doc, """module.exports = {
  apps: [
    {
      name: "zatca-api",
      cwd: "./artifacts/api-server",
      script: "pnpm",
      args: "start",
      env: { NODE_ENV: "production", PORT: 8080 },
      instances: 2,
      exec_mode: "cluster",
      max_memory_restart: "1G",
    },
  ],
};""", "JavaScript")

    add_code_block(doc, """pm2 start ecosystem.config.cjs
pm2 save
pm2 startup    # نفّذ السطر الذي يعرضه ليعمل عند الإقلاع التلقائي""", "Bash")

    add_heading_ar(doc, "الخطوة 9 — إعداد Nginx", level=2)
    add_para_ar(doc, "أنشئ الملف /etc/nginx/sites-available/zatca:")
    add_code_block(doc, """server {
    listen 80;
    server_name yourdomain.sa www.yourdomain.sa;

    # ملفات الواجهة المبنية
    root /home/zatca/zatca/artifacts/zatca-invoicing/dist;
    index index.html;

    client_max_body_size 50M;

    # توجيه الـ API للخادم
    location /api/ {
        proxy_pass http://localhost:8080;
        proxy_http_version 1.1;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
        proxy_read_timeout 300s;
    }

    # نقاط البيع (اختياري)
    location /pos/ {
        alias /home/zatca/zatca/artifacts/pos/dist/;
        try_files $uri $uri/ /pos/index.html;
    }

    # SPA fallback للواجهة الرئيسية
    location / {
        try_files $uri $uri/ /index.html;
    }

    # ضغط الملفات
    gzip on;
    gzip_types text/plain text/css application/json application/javascript text/xml application/xml application/xml+rss text/javascript image/svg+xml;
}""", "nginx.conf")

    add_code_block(doc, """sudo ln -s /etc/nginx/sites-available/zatca /etc/nginx/sites-enabled/
sudo nginx -t
sudo systemctl reload nginx""", "Bash")

    add_heading_ar(doc, "الخطوة 10 — تفعيل HTTPS بشهادة مجانية", level=2)
    add_code_block(doc, """sudo certbot --nginx -d yourdomain.sa -d www.yourdomain.sa
# اتبع التعليمات؛ سيُجدّد تلقائيًا""", "Bash")

    add_heading_ar(doc, "الخطوة 11 — النسخ الاحتياطي الدوري", level=2)
    add_para_ar(doc, "أنشئ سكربت النسخ الاحتياطي /home/zatca/backup.sh:")
    add_code_block(doc, """#!/bin/bash
DATE=$(date +%Y%m%d-%H%M%S)
BACKUP_DIR=/var/backups/zatca
mkdir -p $BACKUP_DIR

# نسخة قاعدة البيانات
PGPASSWORD=YOUR_DB_PASS pg_dump -U zatca_prod_user -h localhost zatca_prod \\
  | gzip > $BACKUP_DIR/db-$DATE.sql.gz

# نسخة ملفات التخزين
tar -czf $BACKUP_DIR/storage-$DATE.tar.gz /var/zatca/storage/

# حذف ما هو أقدم من 30 يومًا
find $BACKUP_DIR -name "*.gz" -mtime +30 -delete""", "Bash")
    add_code_block(doc, """chmod +x /home/zatca/backup.sh
crontab -e
# أضف السطر التالي ليعمل النسخ يوميًا الساعة 3 فجرًا
0 3 * * * /home/zatca/backup.sh""", "Bash")

    doc.add_page_break()

    # ===== 5) Env vars =====
    add_heading_ar(doc, "5. شرح متغيرات البيئة", level=1, color=RGBColor(0x1E, 0x3A, 0x8A))
    add_table_ar(
        doc,
        ["المتغير", "إلزامي؟", "الوصف"],
        [
            ["DATABASE_URL", "نعم", "رابط الاتصال بقاعدة PostgreSQL"],
            ["SESSION_SECRET", "نعم", "مفتاح تشفير الجلسات (≥ 64 حرف عشوائي)"],
            ["DEFAULT_OBJECT_STORAGE_BUCKET_ID", "نعم", "معرّف حاوية تخزين الملفات"],
            ["PRIVATE_OBJECT_DIR", "نعم", "مسار الملفات الخاصة (مرفقات الفواتير)"],
            ["PUBLIC_OBJECT_SEARCH_PATHS", "نعم", "مسار الملفات العامة (شعارات، صور)"],
            ["NODE_ENV", "إنتاج", "production في الإنتاج، development محليًا"],
            ["PORT", "اختياري", "منفذ الخادم (افتراضي 8080)"],
            ["SMTP_HOST/PORT/USER/PASS/FROM", "اختياري", "إعدادات البريد لإرسال التقارير والإشعارات"],
            ["SPL_API_KEY", "اختياري", "مفتاح خدمة العنوان السعودي للتحقق من العناوين"],
            ["TURNSTILE_SECRET_KEY", "اختياري", "حماية النماذج من البوتات (Cloudflare)"],
            ["TURNSTILE_SITE_KEY", "اختياري", "نفس مفتاح Turnstile (للجانب العام)"],
            ["VITE_TURNSTILE_SITE_KEY", "اختياري", "نسخة الواجهة من مفتاح Turnstile"],
        ],
    )

    doc.add_page_break()

    # ===== 6) Updates =====
    add_heading_ar(doc, "6. التحديثات وعمليات الصيانة", level=1, color=RGBColor(0x05, 0x96, 0x69))

    add_heading_ar(doc, "تحديث المشروع", level=2)
    add_code_block(doc, """cd ~/zatca

# سحب آخر التحديثات (لو Git)
git pull

# تركيب أي اعتماديات جديدة
pnpm install --frozen-lockfile

# تحديث مخطط القاعدة (لو فيه تغييرات)
pnpm --filter @workspace/api-server run db:push

# إعادة بناء الواجهات
pnpm --filter @workspace/zatca-invoicing run build
pnpm --filter @workspace/pos run build
pnpm --filter @workspace/api-server run build

# إعادة تشغيل الخادم
pm2 restart zatca-api

# إعادة تحميل Nginx
sudo systemctl reload nginx""", "Bash")

    add_heading_ar(doc, "مراقبة السجلات", level=2)
    add_code_block(doc, """# سجلات PM2 المباشرة
pm2 logs zatca-api

# سجلات Nginx
sudo tail -f /var/log/nginx/access.log
sudo tail -f /var/log/nginx/error.log

# حالة الخدمات
pm2 status
sudo systemctl status nginx postgresql""", "Bash")

    add_heading_ar(doc, "استعادة نسخة احتياطية", level=2)
    add_code_block(doc, """# قاعدة البيانات
gunzip -c /var/backups/zatca/db-YYYYMMDD-HHMMSS.sql.gz \\
  | PGPASSWORD=YOUR_DB_PASS psql -U zatca_prod_user -h localhost zatca_prod

# الملفات
sudo tar -xzf /var/backups/zatca/storage-YYYYMMDD-HHMMSS.tar.gz -C /""", "Bash")

    doc.add_page_break()

    # ===== 7) Troubleshooting =====
    add_heading_ar(doc, "7. حل المشكلات الشائعة", level=1, color=RGBColor(0xDC, 0x26, 0x26))

    add_table_ar(
        doc,
        ["المشكلة", "السبب المحتمل", "الحل"],
        [
            ["الواجهة لا تفتح", "Nginx متوقف أو إعداداته خطأ", "sudo nginx -t ثم sudo systemctl restart nginx"],
            ["خطأ 502 Bad Gateway", "خادم الـ API متوقف", "pm2 restart zatca-api ثم pm2 logs"],
            ["خطأ اتصال قاعدة البيانات", "DATABASE_URL خطأ أو PostgreSQL متوقف", "تأكد من sudo systemctl status postgresql ومن صحة كلمة المرور"],
            ["pnpm install فشل", "ذاكرة غير كافية", "أضف Swap أو زد ذاكرة السيرفر"],
            ["شهادة HTTPS انتهت", "Certbot لم يجدد", "sudo certbot renew --dry-run"],
            ["الملفات المرفوعة لا تظهر", "صلاحيات مجلد التخزين خطأ", "sudo chown -R zatca:zatca /var/zatca/storage"],
            ["السوبر أدمن لا يستطيع الدخول", "لم يُرقَّ الحساب بعد", "نفّذ UPDATE users SET role='superadmin'"],
            ["خطأ توقيع ZATCA", "شهادة المنتج غير مفعّلة", "ادخل صفحة إعدادات ZATCA وفعّل الشهادة عبر OTP من البوابة"],
        ],
    )

    add_heading_ar(doc, "أوامر تشخيص سريعة", level=2)
    add_code_block(doc, """# هل المنفذ مفتوح؟
sudo ss -tlnp | grep -E '80|443|8080|5432'

# مساحة القرص
df -h

# الذاكرة
free -h

# آخر 50 سطر من سجل الخادم
pm2 logs zatca-api --lines 50 --nostream""", "Bash")

    doc.add_page_break()

    # Footer / closing
    add_heading_ar(doc, "ملاحظة ختامية", level=1, color=RGBColor(0x1E, 0x3A, 0x8A))
    add_para_ar(
        doc,
        "هذا الدليل يغطي السيناريو القياسي للتركيب والنشر. لو احتجت دعمًا في إعداد التكامل مع بوابة ZATCA الإنتاجية أو ربط بوابات الدفع أو البريد المؤسسي أو إعداد التوسع الأفقي (load balancer + multiple servers)، تواصل مع فريق الدعم الفني.",
    )
    add_para_ar(doc, "حظًا موفقًا!", bold=True, size=14, color=RGBColor(0x05, 0x96, 0x69))

    out_path = "exports/دليل_تركيب_نظام_زاتكا.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
